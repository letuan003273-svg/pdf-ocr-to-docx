# app.py
import streamlit as st
from pdf2docx import Converter
import img2pdf
import tempfile
import os
from pathlib import Path
import io
import zipfile
from typing import List, Tuple
from PIL import Image
import traceback
import sys
import time

# Optional OCR fallback libraries
try:
    from pdf2image import convert_from_path
    import pytesseract
    from docx import Document
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

st.set_page_config(page_title="PDF/Image → DOCX (OCR & Layout)", layout="centered")

st.title("Chuyển đổi PDF / Ảnh sang Word (.docx) — Giữ layout cố gắng tối đa")
st.markdown(
    """
- Hỗ trợ nhiều file cùng lúc (PDF, PNG, JPG, JPEG).  
- Quy trình ưu tiên: **PDF** → `pdf2docx`. **Ảnh** → convert sang PDF bằng `img2pdf` → `pdf2docx`.  
- Nếu `pdf2docx` thất bại cho file (ví dụ scanned image PDF), app sẽ thử **fallback OCR** (nếu hệ thống có `tesseract` + `poppler`).  
"""
)

uploaded_files = st.file_uploader(
    "Tải lên các file (PDF, PNG, JPG, JPEG). Có thể chọn nhiều file.",
    type=["pdf", "png", "jpg", "jpeg"],
    accept_multiple_files=True
)

zip_output_option = st.checkbox("Nén tất cả kết quả thành 1 file ZIP để tải về", value=True)
process_btn = st.button("Bắt đầu chuyển đổi")

# Helpers
def save_uploadedfile_to_temp(uploaded_file, dirpath) -> Path:
    path = Path(dirpath) / uploaded_file.name
    with open(path, "wb") as f:
        f.write(uploaded_file.read())
    return path

def image_bytes_to_pdf_bytes(image_bytes: bytes) -> bytes:
    # Use img2pdf to convert image bytes into a single-page PDF bytes
    try:
        img = Image.open(io.BytesIO(image_bytes))
        # Ensure RGB for some formats
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        output = img2pdf.convert(img.filename) if getattr(img, "filename", None) else img2pdf.convert(Image.frombytes(img.mode, img.size, img.tobytes()))
        # Above is an attempt but img2pdf.convert needs filenames or PIL images; easier: use BytesIO and pillow save to PDF
    except Exception:
        # fallback: save PIL image to BytesIO as PDF
        img = Image.open(io.BytesIO(image_bytes))
        pdf_io = io.BytesIO()
        img.save(pdf_io, format="PDF", resolution=100.0)
        return pdf_io.getvalue()
    # If we reached here, return output
    return output

def image_file_to_pdf_path(image_path: Path, out_pdf_path: Path) -> Path:
    # Convert image file to a PDF file (single page)
    try:
        with Image.open(image_path) as img:
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(out_pdf_path, "PDF", resolution=100.0)
        return out_pdf_path
    except Exception as e:
        raise

def pdf_to_docx(pdf_path: Path, docx_path: Path) -> None:
    # Use pdf2docx Converter
    conv = None
    try:
        conv = Converter(str(pdf_path))
        conv.convert(str(docx_path), start=0, end=None)
    finally:
        if conv:
            conv.close()

def ocr_pdf_to_docx(pdf_path: Path, docx_path: Path) -> None:
    """
    Fallback OCR: convert PDF pages to images, run Tesseract to extract text, write to a .docx.
    This will NOT preserve layout perfectly but serves as a fallback.
    Requires pdf2image + pytesseract + python-docx + poppler + tesseract installed.
    """
    if not OCR_AVAILABLE:
        raise RuntimeError("OCR stack not available (pdf2image/pytesseract/docx missing).")
    pages = convert_from_path(str(pdf_path))
    doc = Document()
    for page_num, pil_im in enumerate(pages, start=1):
        text = pytesseract.image_to_string(pil_im)
        # Basic: add a heading for page and the text. Preserve blank lines.
        doc.add_paragraph(f"--- Page {page_num} ---")
        for para in text.split("\n\n"):
            cleaned = para.strip()
            if cleaned:
                doc.add_paragraph(cleaned)
        # Insert page break
        if page_num < len(pages):
            doc.add_page_break()
    doc.save(str(docx_path))

def process_single_file(input_path: Path, temp_dir: Path) -> Tuple[bool, str, Path]:
    """
    Returns (success, message, output_docx_path_if_any)
    """
    suffix = input_path.suffix.lower()
    base_name = input_path.stem
    out_docx = temp_dir / f"{base_name}.docx"
    try:
        if suffix == ".pdf":
            # Try direct pdf2docx conversion
            try:
                pdf_to_docx(input_path, out_docx)
                return True, "Converted by pdf2docx", out_docx
            except Exception as e:
                # Attempt fallback OCR if available
                st.warning(f"pdf2docx failed for {input_path.name}: {e}")
                if OCR_AVAILABLE:
                    try:
                        ocr_pdf_to_docx(input_path, out_docx)
                        return True, "Converted by OCR fallback", out_docx
                    except Exception as e2:
                        return False, f"Both pdf2docx and OCR fallback failed: {e2}", None
                else:
                    return False, f"pdf2docx failed and OCR not available: {e}", None

        elif suffix in [".png", ".jpg", ".jpeg"]:
            # Convert image to a temporary PDF, then pdf2docx
            tmp_pdf = temp_dir / f"{base_name}_converted.pdf"
            try:
                image_file_to_pdf_path(input_path, tmp_pdf)
            except Exception as e_img:
                # If direct PIL save fails, try to read bytes and save
                try:
                    with Image.open(str(input_path)) as img:
                        if img.mode in ("RGBA", "P"):
                            img = img.convert("RGB")
                        img.save(tmp_pdf, "PDF", resolution=100.0)
                except Exception as e2:
                    return False, f"Không thể chuyển ảnh sang PDF: {e2}", None

            # Now attempt pdf2docx
            try:
                pdf_to_docx(tmp_pdf, out_docx)
                return True, "Image->PDF->pdf2docx", out_docx
            except Exception as e:
                st.warning(f"pdf2docx failed on image-converted PDF for {input_path.name}: {e}")
                if OCR_AVAILABLE:
                    try:
                        ocr_pdf_to_docx(tmp_pdf, out_docx)
                        return True, "Image->PDF->OCR fallback", out_docx
                    except Exception as e2:
                        return False, f"pdf2docx & OCR fallback failed for image: {e2}", None
                else:
                    return False, f"pdf2docx failed for image and OCR not available: {e}", None
        else:
            return False, f"Unsupported file type: {suffix}", None
    except Exception as e_outer:
        tb = traceback.format_exc()
        return False, f"Unexpected error: {e_outer}\n{tb}", None

# Main processing logic
if process_btn:
    if not uploaded_files or len(uploaded_files) == 0:
        st.info("Vui lòng chọn ít nhất một file để xử lý.")
    else:
        status_container = st.container()
        progress_bar = st.progress(0)
        total = len(uploaded_files)
        processed_results = []  # list of tuples (original_name, success(bool), message, bytes_of_docx)
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            for idx, up in enumerate(uploaded_files, start=1):
                status_slot = status_container.empty()
                status_slot.info(f"[{idx}/{total}] Bắt đầu xử lý: {up.name}")
                # Save uploaded file to temp path
                try:
                    saved_path = save_uploadedfile_to_temp(up, tmpdir)
                except Exception as e:
                    processed_results.append((up.name, False, f"Không lưu được file tạm: {e}", None))
                    progress_bar.progress(int((idx/total)*100))
                    continue

                try:
                    ok, msg, out_docx_path = process_single_file(saved_path, tmpdir_path)
                    if ok and out_docx_path and out_docx_path.exists():
                        # read bytes
                        with open(out_docx_path, "rb") as f:
                            docx_bytes = f.read()
                        processed_results.append((up.name, True, msg, docx_bytes))
                        status_slot.success(f"[{idx}/{total}] Hoàn tất: {up.name} — {msg}")
                    else:
                        processed_results.append((up.name, False, msg, None))
                        status_slot.error(f"[{idx}/{total}] Lỗi: {up.name} — {msg}")
                except Exception as e:
                    tb = traceback.format_exc()
                    processed_results.append((up.name, False, f"Unhandled error: {e}\n{tb}", None))
                    status_slot.error(f"[{idx}/{total}] Lỗi không mong muốn: {up.name} — {e}")

                # update progress
                progress_bar.progress(int((idx/total)*100))
                # small sleep to make UI updates smoother
                time.sleep(0.1)

        # Show summary and download buttons
        st.subheader("Kết quả xử lý")
        success_count = sum(1 for r in processed_results if r[1])
        fail_count = len(processed_results) - success_count
        st.write(f"✅ Thành công: **{success_count}** — ❌ Thất bại: **{fail_count}**")

        # Individual download buttons
        download_entries = []
        for original_name, ok, msg, docx_bytes in processed_results:
            if ok and docx_bytes:
                st.markdown(f"**{original_name}** — {msg}")
                # Use original filename with .docx
                suggested_name = f"{Path(original_name).stem}.docx"
                st.download_button(
                    label=f"Tải xuống {suggested_name}",
                    data=docx_bytes,
                    file_name=suggested_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                download_entries.append((suggested_name, docx_bytes))
            else:
                st.markdown(f"**{original_name}** — ❌ {msg}")

        # Zip option
        if zip_output_option and len(download_entries) > 0:
            zip_io = io.BytesIO()
            with zipfile.ZipFile(zip_io, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                for fname, bts in download_entries:
                    zf.writestr(fname, bts)
            zip_bytes = zip_io.getvalue()
            st.download_button(
                label="Tải về tất cả (ZIP)",
                data=zip_bytes,
                file_name="converted_docs.zip",
                mime="application/zip"
            )

        # If none succeeded, helpful note
        if success_count == 0:
            st.warning(
                "Không có file nào được chuyển đổi thành công. "
                "Nếu các file là ảnh quét (scanned images), app sẽ cần Tesseract + Poppler để chạy chế độ OCR fallback.\n"
                "Vui lòng kiểm tra log lỗi ở trên."
            )

# Footer / help
st.markdown("---")
st.markdown(
    """
**Ghi chú kỹ thuật và troubleshooting**
- Nếu ứng dụng báo lỗi liên quan đến `pdf2docx`, hãy kiểm tra xem file PDF có bị bảo vệ (password-protected) hay không.  
- Để bật OCR fallback (khi `pdf2docx` không giữ được layout do file là ảnh), cần cài đặt thêm `tesseract` và `poppler` trên hệ thống.  
- Thư viện cần thiết xem trong `requirements.txt` bên dưới.
"""
)
